#!/usr/bin/env python3
"""
대한민국 약전 대용량 파일 분할 파싱

대용량 파일(>3MB)을 작은 청크로 분할하여 Upstage API로 파싱
- docx: 문단 기준 분할
- hwpx: 수동 변환 필요 (향후 지원)
"""
import sys
from pathlib import Path
import time
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import tempfile

# shared 모듈 경로 추가
sys.path.insert(0, str(Path(__file__).parent.parent))

from shared.parsers import create_parser


def estimate_docx_size(doc):
    """docx Document 객체의 대략적인 크기 추정 (문단 수 기준)"""
    return len(doc.paragraphs)


def split_docx_by_paragraphs(input_path, chunk_size=100):
    """
    docx 파일을 문단 기준으로 분할

    Args:
        input_path: 입력 docx 파일 경로
        chunk_size: 청크당 문단 수 (기본값: 100)

    Returns:
        List[Path]: 생성된 임시 docx 파일 경로들
    """
    print(f"\n[SPLIT] Reading document: {input_path.name}")
    doc = Document(input_path)

    total_paragraphs = len(doc.paragraphs)
    print(f"  Total paragraphs: {total_paragraphs}")
    print(f"  Chunk size: {chunk_size} paragraphs")

    chunks = []
    temp_dir = Path(tempfile.gettempdir()) / "mfds_split"
    temp_dir.mkdir(exist_ok=True)

    for chunk_idx in range(0, total_paragraphs, chunk_size):
        # 새 Document 생성
        chunk_doc = Document()

        # 문단 복사
        end_idx = min(chunk_idx + chunk_size, total_paragraphs)
        for i in range(chunk_idx, end_idx):
            para = doc.paragraphs[i]
            # 텍스트만 복사 (간단한 방법)
            chunk_doc.add_paragraph(para.text, style=para.style)

        # 임시 파일로 저장
        chunk_num = chunk_idx // chunk_size + 1
        chunk_path = temp_dir / f"{input_path.stem}_chunk_{chunk_num:03d}.docx"
        chunk_doc.save(str(chunk_path))

        chunk_file_size = chunk_path.stat().st_size
        print(f"  Chunk {chunk_num}: paragraphs {chunk_idx+1}-{end_idx} "
              f"({chunk_file_size/1024:.1f} KB)")

        chunks.append(chunk_path)

    print(f"  Created {len(chunks)} chunks")
    return chunks


def parse_large_file(file_info, parser, output_dir, chunk_size=100):
    """
    대용량 파일을 분할하여 파싱

    Args:
        file_info: 파일 정보 dict
        parser: Upstage parser 객체
        output_dir: 출력 디렉토리
        chunk_size: docx 분할 시 청크당 문단 수

    Returns:
        dict: 파싱 결과
    """
    file_path = file_info['path']

    print(f"\n{'='*80}")
    print(f"PARSING LARGE FILE: {file_info['name']}")
    print(f"  Language: {file_info['lang']}")
    print(f"  Format: {file_info['format']}")
    print(f"  Size: {file_path.stat().st_size / 1024 / 1024:.1f} MB")
    print(f"{'='*80}")

    start_time = time.time()

    # 파일 형식에 따라 분할 방법 선택
    if file_info['format'] == 'docx':
        try:
            chunks = split_docx_by_paragraphs(file_path, chunk_size=chunk_size)
        except Exception as e:
            print(f"  [ERROR] Failed to split docx: {e}")
            return {
                'file': file_info['name'],
                'lang': file_info['lang'],
                'format': file_info['format'],
                'status': 'error',
                'error': f"Split failed: {e}"
            }
    elif file_info['format'] == 'hwpx':
        print(f"  [SKIP] hwpx format not yet supported for splitting")
        print(f"  Please convert to PDF manually and use PDF split parser")
        return {
            'file': file_info['name'],
            'lang': file_info['lang'],
            'format': file_info['format'],
            'status': 'skipped',
            'reason': 'hwpx splitting not supported'
        }
    else:
        print(f"  [ERROR] Unsupported format: {file_info['format']}")
        return {
            'file': file_info['name'],
            'lang': file_info['lang'],
            'format': file_info['format'],
            'status': 'error',
            'error': f"Unsupported format: {file_info['format']}"
        }

    # 각 청크 파싱
    print(f"\n[PARSE] Processing {len(chunks)} chunks...")

    all_content = []
    all_html = []
    total_pages = 0
    success_chunks = 0
    error_chunks = 0

    for i, chunk_path in enumerate(chunks, 1):
        print(f"\n  [{i}/{len(chunks)}] Parsing chunk: {chunk_path.name}")
        print(f"    Size: {chunk_path.stat().st_size / 1024:.1f} KB")

        try:
            # 파싱
            chunk_start = time.time()
            result = parser.parse(chunk_path)
            chunk_time = time.time() - chunk_start

            pages = result.get('pages', 0)
            content = result.get('content', '')
            html = result.get('html', '')

            print(f"    [OK] SUCCESS ({chunk_time:.1f}s)")
            print(f"      Pages: {pages}")
            print(f"      Content: {len(content):,} chars")

            all_content.append(content)
            all_html.append(html)
            total_pages += pages
            success_chunks += 1

            # API 부하 방지
            if i < len(chunks):
                time.sleep(1)

        except Exception as e:
            print(f"    [ERROR] {type(e).__name__}: {e}")
            error_chunks += 1
            # 에러가 발생해도 계속 진행
            continue

    # 청크 파일 정리
    print(f"\n[CLEANUP] Removing temporary chunk files...")
    for chunk_path in chunks:
        try:
            chunk_path.unlink()
        except Exception as e:
            print(f"  Warning: Failed to delete {chunk_path}: {e}")

    # 결과 병합
    if success_chunks == 0:
        return {
            'file': file_info['name'],
            'lang': file_info['lang'],
            'format': file_info['format'],
            'status': 'error',
            'error': f'All {len(chunks)} chunks failed to parse'
        }

    merged_content = '\n\n'.join(all_content)
    merged_html = '\n'.join(all_html)

    elapsed = time.time() - start_time

    print(f"\n[MERGE] Combining results...")
    print(f"  Success chunks: {success_chunks}/{len(chunks)}")
    print(f"  Total pages: {total_pages}")
    print(f"  Merged content: {len(merged_content):,} chars")
    print(f"  Merged HTML: {len(merged_html):,} chars")
    print(f"  Total time: {elapsed:.1f}s ({elapsed/60:.1f} min)")

    # 병합된 결과 저장
    output_file = output_dir / f"{file_info['lang']}_{file_info['name']}.json"
    merged_result = {
        'content': merged_content,
        'html': merged_html,
        'pages': total_pages,
        'metadata': {
            'split_parsing': True,
            'total_chunks': len(chunks),
            'success_chunks': success_chunks,
            'error_chunks': error_chunks,
            'chunk_size': chunk_size,
            'parse_time': elapsed
        }
    }

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(merged_result, f, ensure_ascii=False, indent=2)

    print(f"  [SAVED] {output_file}")

    return {
        'file': file_info['name'],
        'lang': file_info['lang'],
        'format': file_info['format'],
        'size_kb': file_path.stat().st_size / 1024,
        'pages': total_pages,
        'chunks': len(chunks),
        'success_chunks': success_chunks,
        'error_chunks': error_chunks,
        'markdown_chars': len(merged_content),
        'html_chars': len(merged_html),
        'parse_time': elapsed,
        'status': 'success' if error_chunks == 0 else 'partial_success',
        'output_file': str(output_file)
    }


def main():
    """대용량 파일 분할 파싱 메인 함수"""

    print("=" * 80)
    print("대한민국 약전 대용량 파일 분할 파싱")
    print("=" * 80)

    # 파싱 실패한 대용량 파일 목록
    base_dir = Path('data/mfds/raw/THE KOREAN PHARMACOPOEIA')

    large_files = [
        # 영문 docx (성공 가능성 높음)
        {
            'path': base_dir / 'en' / '05_[Appendix 4] Monographs_Part2.docx',
            'lang': 'en',
            'format': 'docx',
            'name': '05_[Appendix 4] Monographs_Part2'
        },
        {
            'path': base_dir / 'en' / '06_[Appendix 5] General Tests.docx',
            'lang': 'en',
            'format': 'docx',
            'name': '06_[Appendix 5] General Tests'
        },
        {
            'path': base_dir / 'en' / '07_[Appendix 6] General Information.docx',
            'lang': 'en',
            'format': 'docx',
            'name': '07_[Appendix 6] General Information'
        },
        {
            'path': base_dir / 'en' / '04_[Appendix 3] Monographs_Part1.docx',
            'lang': 'en',
            'format': 'docx',
            'name': '04_[Appendix 3] Monographs_Part1'
        },
        # 한글 hwpx (스킵 예정)
        {
            'path': base_dir / 'ko' / '[별표 4] 의약품각조 제2부 (제2025-18호).hwpx',
            'lang': 'ko',
            'format': 'hwpx',
            'name': '[별표 4] 의약품각조 제2부 (제2025-18호)'
        },
        {
            'path': base_dir / 'ko' / '[별표 5] 일반시험법 (제2025-18호).hwpx',
            'lang': 'ko',
            'format': 'hwpx',
            'name': '[별표 5] 일반시험법 (제2025-18호)'
        },
        {
            'path': base_dir / 'ko' / '[별표 6] 일반정보 (제2025-18호).hwpx',
            'lang': 'ko',
            'format': 'hwpx',
            'name': '[별표 6] 일반정보 (제2025-18호)'
        },
        {
            'path': base_dir / 'ko' / '[별표 3] 의약품각조 제1부 (제2025-18호).hwpx',
            'lang': 'ko',
            'format': 'hwpx',
            'name': '[별표 3] 의약품각조 제1부 (제2025-18호)'
        },
    ]

    # 파일 존재 여부 확인
    print(f"\n[FILES] Checking {len(large_files)} large files...")
    valid_files = []
    for file_info in large_files:
        if file_info['path'].exists():
            size_mb = file_info['path'].stat().st_size / 1024 / 1024
            print(f"  [OK] {file_info['name'][:60]:60s} ({size_mb:.1f} MB)")
            valid_files.append(file_info)
        else:
            print(f"  [SKIP] File not found: {file_info['name']}")

    if not valid_files:
        print("\n[ERROR] No valid files found!")
        return 1

    # 파서 생성
    print(f"\n[PARSER] Creating Upstage parser...")
    try:
        parser = create_parser("upstage")
    except Exception as e:
        print(f"[ERROR] Failed to create parser: {e}")
        return 1

    # 출력 디렉토리
    output_dir = Path('data/mfds/parsed')
    output_dir.mkdir(parents=True, exist_ok=True)

    # 파싱 시작
    results = []
    start_time = time.time()

    for i, file_info in enumerate(valid_files, 1):
        print(f"\n{'='*80}")
        print(f"[{i}/{len(valid_files)}] Processing: {file_info['name']}")
        print(f"{'='*80}")

        result = parse_large_file(
            file_info,
            parser,
            output_dir,
            chunk_size=100  # 청크당 100 문단
        )
        results.append(result)

        # 파일 간 대기
        if i < len(valid_files):
            time.sleep(2)

    # 완료
    elapsed = time.time() - start_time

    print("\n" + "=" * 80)
    print("SUMMARY")
    print("=" * 80)

    success = sum(1 for r in results if r['status'] == 'success')
    partial = sum(1 for r in results if r['status'] == 'partial_success')
    skipped = sum(1 for r in results if r['status'] == 'skipped')
    error = sum(1 for r in results if r['status'] == 'error')

    print(f"\nTotal files: {len(results)}")
    print(f"  Success: {success}")
    print(f"  Partial success: {partial}")
    print(f"  Skipped: {skipped}")
    print(f"  Error: {error}")
    print(f"  Total time: {elapsed:.1f}s ({elapsed/60:.1f} min)")

    # 언어별 통계
    print("\n[BY STATUS]")
    for r in results:
        status_marker = {
            'success': '[OK]',
            'partial_success': '[PARTIAL]',
            'skipped': '[SKIP]',
            'error': '[ERROR]'
        }.get(r['status'], '[?]')

        info = f"{status_marker} {r['file'][:50]:50s} ({r['format']})"
        if r['status'] in ['success', 'partial_success']:
            info += f" - {r['pages']} pages, {r.get('chunks', 0)} chunks"
        elif r['status'] == 'error':
            info += f" - {r.get('error', 'Unknown error')}"
        elif r['status'] == 'skipped':
            info += f" - {r.get('reason', 'Unknown reason')}"

        print(f"  {info}")

    # 결과 저장
    summary_file = output_dir / 'split_parse_summary.json'
    summary = {
        'timestamp': datetime.now().isoformat(),
        'total_files': len(results),
        'success': success,
        'partial_success': partial,
        'skipped': skipped,
        'error': error,
        'elapsed_seconds': elapsed,
        'results': results
    }

    with open(summary_file, 'w', encoding='utf-8') as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\nSummary saved: {summary_file}")

    return 0 if error == 0 else 1


if __name__ == '__main__':
    sys.exit(main())
